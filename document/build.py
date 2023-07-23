import os
from docx.shared import Pt, RGBColor
from constants.globals import DEFAULT_TEMPLATE_FILE_NAME, TEMPLATE_DIR_NAME, TEMPLATE_OUTPUT_DIR_NAME
from docx import Document
import docx

from models.resume import Resume

def create_with_default_template(resume: Resume):
    document = Document(os.path.join(TEMPLATE_DIR_NAME, DEFAULT_TEMPLATE_FILE_NAME))

    document.paragraphs[0].runs[0].text = '     ' + resume.header.intro

    # header table
    header_container_table = document.tables[0]
    header_table = header_container_table.cell(0, 0).tables[0]
    
    # name table
    header_table.cell(0, 0).text = resume.header.name.upper()
    header_table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(30)

    # detail table
    detail_table = header_table.cell(0, 1).tables[0]
    for row in detail_table.rows:
        key = row.cells[1].text
        detail_rows = list(filter(lambda detail: detail.field_icon == key, resume.header.details))
        if len(detail_rows) == 0:
            remove_row(detail_table, row)
            
            continue
        
        row.cells[1].text = ''
        add_hyperlink(row.cells[1].paragraphs[0], detail_rows[0].field_value, detail_rows[0].field_value)

    contentTable = document.tables[1]

    # Academic Profile
    add_academic_profile(resume.educations, contentTable.cell(0, 1).tables[0])
    add_experience(resume.internships, contentTable.cell(2, 1).tables[0])
    add_experience(resume.work_experiences, contentTable.cell(4, 1).tables[0])
    add_projects(resume.projects, contentTable.cell(6, 1).tables[0])
    add_skills(resume.skillset, contentTable.cell(8, 1).tables[0])
    add_achievements(resume.achievements, contentTable.cell(10, 1).tables[0])

    add_array_field(resume.interests, contentTable.cell(12, 1).tables[0])
    add_array_field(resume.hobbies, contentTable.cell(14, 1).tables[0])
    add_array_field(resume.languages, contentTable.cell(16, 1).tables[0])
    

    
    document.save(os.path.join(TEMPLATE_OUTPUT_DIR_NAME, resume.header.name + ' Resume' + ".docx"))


def add_academic_profile(educations, academic_profile_table):
    for row in academic_profile_table.rows:
        remove_row(academic_profile_table, row)

    if len(educations) == 0:
        return
    
    for education in educations:
        academic_profile_table.add_row()
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].text = education.institute
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].paragraphs[0].runs[0].underline  = True
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(11)

        academic_profile_table.add_row()
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].text = education.degree

        if education.branch != "":
            academic_profile_table.add_row()
            academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].text = education.branch
        
        academic_profile_table.add_row()
        if education.duration.from_date != '':
            academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].text = education.duration.from_date + " - " + education.duration.to_date
        else:
            academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].text = education.duration.to_date

        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        if education.duration.duration != '':
            academic_profile_table.add_row()
            academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].text = education.duration.duration

        
        academic_profile_table.add_row()
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].text = education.score
        academic_profile_table.rows[len(academic_profile_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        academic_profile_table.add_row()

    remove_row(academic_profile_table, academic_profile_table.rows[len(academic_profile_table.rows)-1])

def add_experience(experiences, experience_table):
    for row in experience_table.rows:
        remove_row(experience_table, row)

    if len(experiences) == 0:
        return

    for experience in experiences:
        experience_table.add_row()
        experience_table.rows[len(experience_table.rows)-1].cells[0].text = experience.organization
        experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].underline  = True
        experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(11)

        experience_table.add_row()
        experience_table.rows[len(experience_table.rows)-1].cells[0].text = experience.title

        if len(experience.description) > 0:
            for description in experience.description:
                experience_table.add_row()
                experience_table.rows[len(experience_table.rows)-1].cells[0].text = ('   •   ' + description)
                experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
                experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(9)

        if len(experience.tech_stack) > 0:
            experience_table.rows[len(experience_table.rows)-1].cells[0].text = ', '.join(experience.tech_stack)
        
        experience_table.add_row()
        if experience.duration.from_date != '':
            experience_table.rows[len(experience_table.rows)-1].cells[0].text = experience.duration.from_date + " - " + experience.duration.to_date
        else:
            experience_table.rows[len(experience_table.rows)-1].cells[0].text = experience.duration.to_date

        experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        if experience.duration.duration != '':
            experience_table.add_row()
            experience_table.rows[len(experience_table.rows)-1].cells[0].text = experience.duration.duration
            experience_table.rows[len(experience_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True

        experience_table.add_row()

    remove_row(experience_table, experience_table.rows[len(experience_table.rows)-1])

def add_projects(projects, projects_table):
    for row in projects_table.rows:
        remove_row(projects_table, row)
    
    if len(projects) == 0:
        return

    for project in projects:
        projects_table.add_row()
        projects_table.rows[len(projects_table.rows)-1].cells[0].text = project.title
        projects_table.rows[len(projects_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        projects_table.rows[len(projects_table.rows)-1].cells[0].paragraphs[0].runs[0].underline  = True
        projects_table.rows[len(projects_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(11)

        if project.motivation != '':
            projects_table.add_row()
            projects_table.rows[len(projects_table.rows)-1].cells[0].text = project.motivation

        if len(project.description) > 0:
            for my_role in project.description:
                projects_table.add_row()
                projects_table.rows[len(projects_table.rows)-1].cells[0].text = ('   •   ' + my_role)
                projects_table.rows[len(projects_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
                projects_table.rows[len(projects_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(9)

        if len(project.tech_stack) > 0:
            projects_table.add_row()
            projects_table.rows[len(projects_table.rows)-1].cells[0].text = ', '.join(project.tech_stack)

        if len(project.other_collaborators) > 0:
            projects_table.add_row()
            collaborators = ', '.join(project.other_collaborators)
            projects_table.rows[len(projects_table.rows)-1].cells[0].text = 'Collaborators: ' + collaborators 

        if len(project.description) > 0:
            projects_table.add_row()
            projects_table.rows[len(projects_table.rows)-1].cells[0].text = "My role: "
            for my_role in project.my_role:
                projects_table.add_row()
                projects_table.rows[len(projects_table.rows)-1].cells[0].text = ('   •   ' + my_role)
                projects_table.rows[len(projects_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
                projects_table.rows[len(projects_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(9)
        
        projects_table.add_row()

    remove_row(projects_table, projects_table.rows[len(projects_table.rows)-1])

def add_skills(skills, skills_table):
    for row in skills_table.rows:
        remove_row(skills_table, row)
    
    if len(skills) == 0:
        return

    for skill in skills:
        skills_table.add_row()
        skills_table.rows[len(skills_table.rows)-1].cells[0].text = skill.skill_name
        skills_table.rows[len(skills_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        skills_table.rows[len(skills_table.rows)-1].cells[0].paragraphs[0].runs[0].underline  = True
        skills_table.rows[len(skills_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(11)

        if len(skill.description) > 0:
            for description in skill.description:
                skills_table.add_row()
                skills_table.rows[len(skills_table.rows)-1].cells[0].text = ('   •   ' + description)
                skills_table.rows[len(skills_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
                skills_table.rows[len(skills_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(9)
        
        skills_table.add_row()

    remove_row(skills_table, skills_table.rows[len(skills_table.rows)-1])

def add_achievements(achievements, achievement_table):
    for row in achievement_table.rows:
        remove_row(achievement_table, row)
    
    if len(achievements) == 0:
        return

    for achievement in achievements:
        achievement_table.add_row()
        achievement_table.rows[len(achievement_table.rows)-1].cells[0].text = achievement.name
        achievement_table.rows[len(achievement_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        achievement_table.rows[len(achievement_table.rows)-1].cells[0].paragraphs[0].runs[0].underline  = True
        achievement_table.rows[len(achievement_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(11)

        if len(achievement.description) > 0:
            for description in achievement.description:
                achievement_table.add_row()
                achievement_table.rows[len(achievement_table.rows)-1].cells[0].text = ('   •   ' + description)
                achievement_table.rows[len(achievement_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
                achievement_table.rows[len(achievement_table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(9)
        
        if achievement.when != '':
            achievement_table.add_row()
            achievement_table.rows[len(achievement_table.rows)-1].cells[0].text = achievement.when
            achievement_table.rows[len(achievement_table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
            achievement_table.rows[len(achievement_table.rows)-1].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        achievement_table.add_row()

    remove_row(achievement_table, achievement_table.rows[len(achievement_table.rows)-1])

def add_array_field(array_field, table):
    for row in table.rows:
        remove_row(table, row)
        
    if len(array_field) == 0:
        return

    for field_value in array_field:
        table.add_row()
        table.rows[len(table.rows)-1].cells[0].text = ('   •   ' + field_value)
        table.rows[len(table.rows)-1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[len(table.rows)-1].cells[0].paragraphs[0].runs[0].font.size  = Pt(10)


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

